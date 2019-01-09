import pandas as pd
import re
from docx.api import Document
from unidecode import unidecode
import numpy as np


def ConvertNumeric(string):
    newString = ''
    for s in string:
        if s.isdigit():
            s = unidecode(s)
        newString += s
    return newString


document = Document('electzone.docx')

data = []
for j in range(len(document.tables)):
    table = document.tables[j]
    keys = None
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)
        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)

df = pd.DataFrame(data)
df = df.applymap(ConvertNumeric)
df = df.apply(lambda x: x.str.strip())
df = df.rename(columns=lambda x: re.sub('\n| ', '', x))

for i in df['ลำดับที่'].index:
    if df['ลำดับที่'].iloc[i]:
        val = df['ลำดับที่'].iloc[i]
    else:
        df['ลำดับที่'].iloc[i] = val

checkdf = df['ท้องที่ที่ประกอบเป็นเขตเลือกตั้ง'][
    df['ท้องที่ที่ประกอบเป็นเขตเลือกตั้ง'].str.contains('\n|[ก-๙]\s', regex=True)]
n = 0
for i in checkdf.index:
    i += n
    df = pd.concat(
        [df.loc[:i],
         pd.DataFrame([[''] * len(df.columns)], columns=df.columns, index=[i + 1]),
         df.loc[i + 1:]]) \
        .reset_index(drop=True)
    splitlist = re.findall('[ก-๙()]+', checkdf.loc[i - n])
    df.loc[i:i + len(splitlist) - 1, 'ท้องที่ที่ประกอบเป็นเขตเลือกตั้ง'] = splitlist
    df.loc[i + 1, ['ลำดับที่', 'เขตเลือกตั้งที่']] = df.loc[i, ['ลำดับที่', 'เขตเลือกตั้งที่']]
    n += 1

for i in range(1, 78):
    i = str(i)
    containcheck = df['เขตเลือกตั้งที่'].loc[df['ลำดับที่'] == i].str.contains('\n')
    vals = df['เขตเลือกตั้งที่'].loc[df['ลำดับที่'] == i][containcheck].unique()
    for val in vals:
        dffind = df['เขตเลือกตั้งที่'][(df['เขตเลือกตั้งที่'] == val) & (df['ลำดับที่'] == i)]
        df['เขตเลือกตั้งที่'][(df['เขตเลือกตั้งที่'] == val) & (df['ลำดับที่'] == i)] = pd.Series(
            list(val) + [val[-1]] * (len(dffind) - len(val)), index=dffind.index)


def dup_blank_row(dfcol):
    for i in df.index:
        if dfcol.iloc[i] in ['\n', '', '.']:
            dfcol.iloc[i] = dfcol.iloc[i - 1]


dup_blank_row(df['เขตเลือกตั้งที่'])
dup_blank_row(df['จังหวัด'])

havecondition = False
for i in df.index:
    zonename = df['ท้องที่ที่ประกอบเป็นเขตเลือกตั้ง'].iloc[i]
    if '(' in zonename:
        havecondition = True
        zonenum = df['เขตเลือกตั้งที่'].iloc[i]
    if havecondition:
        df['เขตเลือกตั้งที่'].iloc[i] = zonenum
    if havecondition or re.match('\(.*|^[2-9].\s.*|.*\)', df['ท้องที่ที่ประกอบเป็นเขตเลือกตั้ง'].iloc[i]):
        df['เขตเลือกตั้งที่'].iloc[i] = df['เขตเลือกตั้งที่'].iloc[i - 1]
    if ')' in zonename:
        havecondition = False


def get_amphor_index(string):
    return df['เขตเลือกตั้งที่'][df['ท้องที่ที่ประกอบเป็นเขตเลือกตั้ง'].str.contains(string)].index[0]


df.loc[df['ท้องที่ที่ประกอบเป็นเขตเลือกตั้ง'] == 'อำเภอเวียงสา', 'เขตเลือกตั้งที่'] = '2'
df.loc[df['ท้องที่ที่ประกอบเป็นเขตเลือกตั้ง'] == 'อำเภอปัว', 'เขตเลือกตั้งที่'] = '3'
df.loc[df['ท้องที่ที่ประกอบเป็นเขตเลือกตั้ง'] == 'อำเภอเชียงคำ', 'เขตเลือกตั้งที่'] = '2'
df.loc[df['ท้องที่ที่ประกอบเป็นเขตเลือกตั้ง'] == 'อำเภอดอกคำใต้', 'เขตเลือกตั้งที่'] = '3'
df.loc[df['ท้องที่ที่ประกอบเป็นเขตเลือกตั้ง'] == 'อำเภอภูกามยาว', 'เขตเลือกตั้งที่'] = '3'
df.loc[df['ท้องที่ที่ประกอบเป็นเขตเลือกตั้ง'] == 'อำเภอกาบัง', 'เขตเลือกตั้งที่'] = '2'
# df['เขตเลือกตั้งที่'].loc[989:997] = '2'
startindex = get_amphor_index('อำเภอโพธาราม')
df['เขตเลือกตั้งที่'].loc[startindex:startindex + 4] = '3'
df['เขตเลือกตั้งที่'].loc[startindex + 5] = '4'

startindex = get_amphor_index('อำเภอโพธิ์ชัย')
df['เขตเลือกตั้งที่'].loc[startindex:startindex + 12] = '2'
# df['เขตเลือกตั้งที่'].loc[1321:1323] = '6'
startindex = get_amphor_index('อำเภอเขวาสินรินทร์')
df['เขตเลือกตั้งที่'].loc[startindex:startindex + 13] = '2'
df['เขตเลือกตั้งที่'].loc[startindex + 32:startindex + 44] = '4'
# df['เขตเลือกตั้งที่'].loc[1498:1501] = '3'
# df['เขตเลือกตั้งที่'].loc[1502:1503] = '4'


df['ท้องที่ที่ประกอบเป็นเขตเลือกตั้ง'].replace('\d{1,2}.\s|^และ', '', regex=True, inplace=True)
df.replace('^ตำบล', '', regex=True, inplace=True)
df.replace('อำเภอ|^ตำบล|เขต|แขวง', '', regex=True, inplace=True)  ### for test


# for i in range(1, 78):
#     i = str(i)
#     containcheck = df['เขตเลือกตั้งที่'].loc[df['ลำดับที่'] == i]
#     vals = df['เขตเลือกตั้งที่'].loc[df['ลำดับที่'] == i].unique()
#     print(f'{i}\n======\n' + df['จังหวัด'].loc[df['ลำดับที่'] == i].iloc[0])
#     for val in vals:
#         print(df['ท้องที่ที่ประกอบเป็นเขตเลือกตั้ง'][(df['เขตเลือกตั้งที่'] == val) & (df['ลำดับที่'] == i)].iloc[0])

def clean_tesban(df):
    # for i in df[~df.str.startswith('ตำบล') & df.str.contains("^\w", regex=True)].index:
    #     if not any(df.iloc[i+1].startswith(x) for x in ['เทศบาล', 'แขวง']) and \
    #             len(df.iloc[i+1])>0:
    #         df.iloc[i]=df.iloc[i]+df.iloc[i+1]
    #         df.iloc[i+1]=''
    for i in df[df.str.startswith('เทศบาลเมือง') | df.str.startswith('เทศบาลนคร')].index:
        if not any(df.iloc[i + 1].startswith(x) for x in ['เทศบาล', 'ตำบล']) and \
                len(df.iloc[i + 1]) > 0 or \
                df.iloc[i].endswith("-") or \
                (df.iloc[i].endswith("นิคมสร้างตนเอง") and df.iloc[i+1].endswith("ลำโดมน้อย")):
            df.iloc[i] = df.iloc[i] + df.iloc[i + 1]
            df.iloc[i + 1] = ''
        df.iloc[i] = df.iloc[i].replace('(', '')
    isstart = False
    for i in df.index:
        if df.iloc[i]:
            isstart = True
            startloc = i - 1
        else:
            startloc = 0
        conditionlist = []
        k = 0
        while isstart:
            if df.iloc[i + k]:
                conditionlist.append(df.iloc[i + k])
            else:
                isstart = False
            df.iloc[i + k] = ''
            k += 1
        if startloc != 0:
            df.iloc[startloc] = conditionlist


def getGeoCode(P, A=None, T=None):
    try:
        if P and (A == None or A == "ทั้งจังหวัด") and T == None:
            return str(shpdf.loc[(shpdf['P_NAME_T'].str.contains(P)), 'P_CODE'].iloc[0])
        if P and A and T == None:
            if A == "ทั้งจังหวัด":
                return None
            else:
                return str(shpdf.loc[(shpdf['P_NAME_T'].str.contains(P)) &
                                     (shpdf['A_NAME_T'].str.contains(A)), 'A_CODE'].iloc[0])
        if P and A and T:
            return str(shpdf.loc[(shpdf['P_NAME_T'].str.contains(P)) &
                                 (shpdf['A_NAME_T'].str.contains(A)) &
                                 (shpdf['T_NAME_T'].str.contains(T)), 'T_CODE'].iloc[0])
    except:
        return None


def clean_brace(dframe, delstart):
    havecondition = False
    findbrace = 0
    for i in df.index:
        zonename = df['ท้องที่ที่ประกอบเป็นเขตเลือกตั้ง'].iloc[i]
        if '(' in zonename:
            findbrace += 1
        if any(('(' + x in zonename) for x in delstart):
            havecondition = True
        if havecondition:
            dframe.iloc[i] = zonename
            df['ท้องที่ที่ประกอบเป็นเขตเลือกตั้ง'].iloc[i] = ''
        if ')' in zonename:
            findbrace -= 1
        if ')' in zonename and findbrace == 0:
            havecondition = False
    return dframe.str.replace(f'(\({"|".join(delstart)}|^และ|\)$)', '', regex=True)

def editTambonDf(P,A,T,new):
    for col in ['interior', 'exterior']:
        try:
            templist = df.loc[(df['จังหวัด'] == P) &
                              (df['ท้องที่ที่ประกอบเป็นเขตเลือกตั้ง'] == A) &
                              (df['interior'].apply(lambda x: T in x)), [col]]
            df.loc[templist.index[0], col] = [new if x == T else x for x in templist.iloc[0][0]]
        except:pass

df['interior'] = ''
df['interior'] = clean_brace(df['interior'], ['เฉพาะตำบล', 'เฉพาะ'])
# df['interior']=clean_brace(df['interior'], 'ใน')

df['exterior'] = ''
df['exterior'] = clean_brace(df['exterior'], ['ยกเว้นตำบล', 'ยกเว้น'])



df.replace('[()]', '', regex=True, inplace=True)
clean_tesban(df.interior)
clean_tesban(df.exterior)

df = df[df['ท้องที่ที่ประกอบเป็นเขตเลือกตั้ง'].map(len) > 0].reset_index(drop=True)

df['ท้องที่ที่ประกอบเป็นเขตเลือกตั้ง'] = df['ท้องที่ที่ประกอบเป็นเขตเลือกตั้ง'].str.strip()
df['ท้องที่ที่ประกอบเป็นเขตเลือกตั้ง'][df['ท้องที่ที่ประกอบเป็นเขตเลือกตั้ง'] == 'กันทรลักษณ์'] = 'กันทรลักษ์'
df['ท้องที่ที่ประกอบเป็นเขตเลือกตั้ง'][df['ท้องที่ที่ประกอบเป็นเขตเลือกตั้ง'] == 'สนามชัย'] = 'สนามชัยเขต'
editTambonDf('สุรินทร์','เมืองสุรินทร์','อ็อง',new='ตาอ็อง')
editTambonDf('สุรินทร์','เมืองสุรินทร์','ปะทัดบุ',new='ประทัดบุ')
editTambonDf('สุรินทร์','เมืองสุรินทร์','ศีขรภูมิ',new='จารพัต')

startindex = df[(df['ลำดับที่'] == '58') & df.ท้องที่ที่ประกอบเป็นเขตเลือกตั้ง.str.contains('เมืองนครสวรรค์')].index
df['interior'].loc[startindex[0]] = ["นครสวรรค์ตก(ในเทศบาลนครนครสวรรค์)",
                                     "นครสวรรค์ออก(ในเทศบาลนครนครสวรรค์)",
                                     "แควใหญ่(ในเทศบาลนครนครสวรรค์)",
                                     "ปากน้ำโพ",
                                     "วัดไทร",
                                     "บางม่วง",
                                     "บ้านมะเกลือ",
                                     "บ้านแก่ง",
                                     "หนองกรด",
                                     "หนองกระโดน",
                                     "บึงเสนาท"]
df['exterior'].loc[startindex[1]] = df['interior'].loc[startindex[0]]
# df['ท้องที่ที่ประกอบเป็นเขตเลือกตั้ง'][df['ท้องที่ที่ประกอบเป็นเขตเลือกตั้ง']=='ป้อมปราบศัตรูพ่า'] = 'ป้อมปราบศัตรูพ่า'

df.to_csv('electzone-processed.csv')

from shptocsv import shptodf

shpdf = shptodf('tambon/TH_Tambon.shp')

shpdf["P_CODE"]=shpdf["P_CODE"].apply(str)
shpdf["A_CODE"]=shpdf["A_CODE"].apply(str)

shpdf['A_NAME_T'] = shpdf['A_NAME_T'].str.strip()
shpdf['A_NAME_T'] = shpdf['A_NAME_T'].str.replace('กิ่งอำเภอ','')
shpdf.loc[shpdf['A_NAME_T'] == 'ป้อมปราบศัตรูพ่า', 'A_NAME_T'] = 'ป้อมปราบศัตรูพ่าย'
shpdf.loc[shpdf['A_NAME_T'] == 'เมืองนครศรีธรรมร', 'A_NAME_T'] = 'เมืองนครศรีธรรมราช'
shpdf.loc[shpdf['A_NAME_T'] == 'เมืองสุราษฎร์ธาน', 'A_NAME_T'] = 'เมืองสุราษฎร์ธานี'
shpdf.loc[shpdf['A_NAME_T'] == 'กรงปีนัง', 'A_NAME_T'] = 'กรงปินัง'
shpdf.loc[shpdf['A_NAME_T'] == 'ป่าพยอม', 'A_NAME_T'] = 'ป่าพะยอม'
shpdf.loc[shpdf['A_NAME_T'] == 'เมืองประจวบคีรีข', 'A_NAME_T'] = 'เมืองประจวบคีรีขันธ์'
shpdf.loc[shpdf['A_NAME_T'] == 'ว่านใหญ่', 'A_NAME_T'] = 'หว้านใหญ่'
shpdf.loc[shpdf['A_NAME_T'] == 'บึงกาฬ', 'A_NAME_T'] = 'เมืองบึงกาฬ'
shpdf.loc[shpdf['A_NAME_T'] == 'สุไหงโกลก', 'A_NAME_T'] = 'สุไหงโก-ลก'

### ตำบลนิคมลำโดมน้อย เผลี่ยนชื่อเป็น นิคมสร้างตนเองลำโดมน้อย
shpdf.loc[shpdf['T_NAME_T'] == 'นิคมลำโดมน้อย', 'T_NAME_T'] = 'นิคมสร้างตนเองลำโดมน้อย'
## เปลี่ยนชื่ออำเภอจาก หนองบุนนาก เป็น หนองบุญมาก
shpdf.loc[shpdf['A_NAME_T'] == 'หนองบุนนาก', 'A_NAME_T'] = 'หนองบุญมาก'

## เปลี่ยน shpdf.loc[shpdf['A_NAME_T'] == 'กิ่งอำเภอเอราวัณ', 'A_NAME_T'] = 'เอราวัณ'
shpdf.loc[shpdf['A_NAME_T'] == 'กิ่งอำเภอเอราวัณ', 'A_NAME_T'] = 'เอราวัณ'

### อำเภอกัลยาณิวัฒนา เชียงใหม่ เปลี่ยนชื่อจาก ตำบลบ้านจันทร์
shpdf.loc[(shpdf['P_NAME_T'] == 'เชียงใหม่') &
           shpdf['A_NAME_T'].str.contains('แม่แจ่ม') &
           shpdf['T_NAME_T'].str.contains('บ้านจันทร์'),
         ['A_NAME_T']] = "กัลยาณิวัฒนา"
shpdf.loc[shpdf['A_NAME_T'] == 'กัลยาณิวัฒนา','A_CODE']="25"

### อำเภแเวียงเก่า ขอนแก่น
shpdf.loc[(shpdf['P_NAME_T'] == 'ขอนแก่น') &
           shpdf['A_NAME_T'].str.contains('ภูเวียง') &
          (shpdf['T_NAME_T'].str.contains('ในเมือง') |
           shpdf['T_NAME_T'].str.contains('เมืองเก่าพัฒนา') |
           shpdf['T_NAME_T'].str.contains('เขาน้อย')),
         ['A_NAME_T']] = "เวียงเก่า"
shpdf.loc[shpdf['A_NAME_T'] == 'เวียงเก่า','A_CODE']="29"




df["p_code"]=df.apply(lambda x: getGeoCode(x['จังหวัด']), axis=1)
df["a_code"]=df.apply(lambda x: getGeoCode(x['จังหวัด'], x['ท้องที่ที่ประกอบเป็นเขตเลือกตั้ง']), axis=1)



# shpdf['A_NAME_T']=shpdf['A_NAME_T'].str.strip()
merge = pd.merge(df, shpdf, how='left', left_on=['p_code', 'a_code'],
                 right_on=['P_CODE', 'A_CODE'])
merrgefail = merge[['จังหวัด', 'ท้องที่ที่ประกอบเป็นเขตเลือกตั้ง']][merge['T_NAME_T'].isna()]

shpdf['A_NAME_T'][(shpdf['P_NAME_T'] == 'ยะลา') & shpdf['A_NAME_T'].str.contains('สนาม')]  ##fortest

# subdistbkkshp = shptodf('subdist_bma/subdist_bma.shp')
# subdistbkkshp["SNAME"] = subdistbkkshp["SNAME"].replace("^แขวง", "", regex=True)
# subdistbkkshp["DNAME"] = subdistbkkshp["DNAME"].replace("^เขต", "", regex=True)

onlydf=pd.DataFrame(pd.DataFrame(df.interior[df.interior != ''].tolist(), index=df.index[df.interior != '']).stack(),columns=['interior'])
tempdf=[]
for i in onlydf.index:
    tempdf.append([df.loc[i[0],["จังหวัด"]].iloc[0],df.loc[i[0],["ท้องที่ที่ประกอบเป็นเขตเลือกตั้ง"]].iloc[0],onlydf.loc[i].iloc[0]])
onlydf=pd.DataFrame(tempdf,columns=["P","A","T"])
onlymerge = pd.merge(onlydf, shpdf, how='left', left_on=['P', 'A', 'T'],
                 right_on=['P_NAME_T', 'A_NAME_T', 'T_NAME_T'])
onlymerrgefail = onlymerge[['P', 'A', 'T']][onlymerge['A_NAME_T'].isna()]


[{"id": "6001",
  "name": "เมืองนครสวรรค์",
  "condition": [{"operator":"interior",
                 "id": "600106",
                "name": "นครสวรรค์ตก",
                "divisonType": "ตำบล"},
                {"divisonType": "เทศบาล",
                 "name": "นครนครสวรรค์",
                 "conditionType": "interior"}]}]

th_map_df=pd.read_csv("th_map.csv")
th_map_df['id'] = th_map_df['id'].apply(str)



import requests

def getShapeFromName(location,wikimapiakey, gmaptoken):
    gmap_client = googlemaps.Client(gmaptoken)
    gmap_result = gmap_client.geocode('เทศบาลเมืองนครพนม')
    point = gmap_result[0]['geometry']['location']
    bbox={'lon_min':f'{point["lng"] - 0.1}',
    'lon_max':f'{point["lng"] + 0.1}',
    'lat_min':f'{point["lat"] - 0.1}',
    'lat_max':f'{point["lat"] + 0.1}'}
    bbox='{0[lon_min]},{0[lat_min]},{0[lon_max]},{0[lat_max]}'.format(bbox)
    # "http://api.wikimapia.org/?key={}&function=box&coordsby=bbox&bbox=104.297,17.042,104.848,17.718&category=7227&count=100&format=json
    url = f"http://api.wikimapia.org/?key={wikimapiakey}&function=box&coordsby=bbox&bbox={bbox}&language=th&category=7227&count=100&format=json"
    j=requests.get(url).json()
    k=[x for x in j['folder'] if x['name']==location][0]
    geom={}
    geom['type']='Polygon'
    geom['coordinates']=[[[x['x'],x['y']] for x in k["polygon"]]]
    return geom

wikimapiakey = ""
gmaptoken = ""
getShapeFromName(location, wikimapiakey, gmaptoken)
import geopandas as gpd
gpd.read_file()

import googlemaps
gmap_client=googlemaps.Client("AIzaSyBkdiNMQbWx9jV4LeM1ulKs7Ejankpnb_I")
gmap_result=gmap_client.geocode('เทศบาลเมืองนครพนม')